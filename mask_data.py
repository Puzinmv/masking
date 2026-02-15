import argparse
import json
import logging
import os
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Iterable, Optional

import yaml
from presidio_analyzer import AnalyzerEngine, RecognizerRegistry, PatternRecognizer, Pattern, RecognizerResult
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_analyzer.nlp_engine import NlpEngine

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import textwrap
from io import StringIO


LOGGER = logging.getLogger("mask_data")


IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tif",
    ".tiff",
    ".gif",
    ".webp",
    ".heic",
    ".heif",
    ".jp2",
    ".j2k",
    ".jpx",
    ".jpf",
    ".jpm",
    ".jxr",
    ".ppm",
    ".pgm",
    ".pbm",
    ".pnm",
    ".ico",
}
SUPPORTED_EXTENSIONS = {".txt", ".docx", ".doc", ".pdf"} | IMAGE_EXTENSIONS

OCR_LANGUAGE = "ru"
OCR_PDF_ZOOM = 2.0
OCR_MIN_CONFIDENCE = 0.3
_OCR_ENGINE = None
_OCR_DEPS = None

# Глобальный словарь "стоп-слов" по типам сущностей, загружается из YAML.
# Ключ: тип сущности (PERSON, ORGANIZATION, ...), значение: набор точных строк и regex-паттернов.


@dataclass(frozen=True)
class StopwordSet:
    exact: set[str]
    regex: List[re.Pattern]


ENTITY_STOPWORDS: Dict[str, StopwordSet] = {}

# Кеш лемм и NLP-движок для морфологической фильтрации стоп-слов.
STOPWORD_LEMMA_CACHE: Dict[str, set[str]] = {}
STOPWORDS_NLP_ENGINE: Optional[NlpEngine] = None
STOPWORDS_LANGUAGE: str = "ru"


@dataclass
class MaskingConfig:
    language: str
    entity_types: List[str]
    score_threshold: float


def load_config(path: Optional[Path]) -> MaskingConfig:
    """
    Загрузка конфигурации из YAML. При отсутствии файла используются значения по умолчанию.
    """
    default = {
        "language": "ru",
        "entity_types": [
            "PERSON",
            "PHONE_NUMBER",
            "EMAIL_ADDRESS",
            "ORGANIZATION",
            "LOCATION",
            "INN",
            "SNILS",
            "OGRN",
            "BANK_ACCOUNT",
            "ORG_INN",
            "ORG_KPP",
            "ORG_OGRN",
        ],
        "score_threshold": 0.35,
    }

    if path is None or not path.is_file():
        cfg = default
    else:
        with path.open("r", encoding="utf-8") as f:
            file_cfg = yaml.safe_load(f) or {}
        cfg = {**default, **file_cfg}

    return MaskingConfig(
        language=cfg["language"],
        entity_types=list(cfg["entity_types"]),
        score_threshold=float(cfg["score_threshold"]),
    )


def load_stopwords(path: Optional[Path]) -> Dict[str, StopwordSet]:
    """
    Загрузка стоп-слов из YAML-файла.
    Формат:
      ORGANIZATION:
        - "СОГЛАСОВАНО"
        - "УТВЕРЖДАЮ"
      PERSON:
        - "КАТЕГОРИЯ"
        - "КАТЕГОРИЯ III"
    Все ключи и значения приводятся к верхнему регистру.
    """
    if path is None:
        # По умолчанию пробуем файл stopwords.yaml в рабочей директории.
        path = Path("stopwords.yaml").resolve()

    if not path.is_file():
        LOGGER.info("Файл стоп-слов не найден: %s (будет использован пустой список)", path)
        return {}

    try:
        with path.open("r", encoding="utf-8") as f:
            raw = yaml.safe_load(f) or {}
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Ошибка чтения файла стоп-слов %s: %s", path, exc)
        return {}

    if not isinstance(raw, dict):
        LOGGER.error("Некорректный формат stopwords YAML (ожидается mapping): %s", path)
        return {}

    result: Dict[str, StopwordSet] = {}
    for ent_type, values in raw.items():
        if not isinstance(values, list):
            continue
        et = str(ent_type).upper()
        exact: set[str] = set()
        regex: List[re.Pattern] = []
        for v in values:
            s = str(v).strip()
            if not s:
                continue
            if s.lower().startswith("re:"):
                pattern = s[3:].strip()
                if not pattern:
                    continue
                try:
                    regex.append(re.compile(pattern, flags=re.IGNORECASE))
                except re.error as exc:
                    LOGGER.warning("Некорректный regex в stopwords (%s): %s", et, exc)
                continue
            exact.add(s.upper())
        if exact or regex:
            result[et] = StopwordSet(exact=exact, regex=regex)

    LOGGER.info("Загружено стоп-слов: %d типов из %s", len(result), path)
    return result


def _get_lemmas_for_text(text: str) -> set[str]:
    """
    Получить множество лемм для строки, используя тот же NLP-движок, что и Presidio.
    Леммы используются для того, чтобы одно стоп-слово (например, "ПОЛОЖЕНИЕ")
    отбрасывало все его склонения ("ПОЛОЖЕНИЯ", "ПОЛОЖЕНИИ" и т.п.).
    """
    original = text.strip()
    if not original:
        return set()

    cache_key = original.upper()

    cached = STOPWORD_LEMMA_CACHE.get(cache_key)
    if cached is not None:
        return cached

    if STOPWORDS_NLP_ENGINE is None:
        # NLP-движок ещё не инициализирован или недоступен —
        # возвращаем пустое множество, чтобы не ломать основной поток.
        lemmas: set[str] = set()
        STOPWORD_LEMMA_CACHE[cache_key] = lemmas
        return lemmas

    try:
        # Для корректной морфологии ру‑модель лучше работает с нижним регистром.
        doc = STOPWORDS_NLP_ENGINE.process_text(original.lower(), STOPWORDS_LANGUAGE)
    except Exception:  # noqa: BLE001
        lemmas = set()
        STOPWORD_LEMMA_CACHE[cache_key] = lemmas
        return lemmas

    # process_text() в Presidio возвращает NlpArtifacts, а не сам spaCy-doc.
    # Используем его поля lemmas/tokens.
    lemmas: set[str] = set()
    doc_lemmas = getattr(doc, "lemmas", None)
    doc_tokens = getattr(doc, "tokens", None)

    if doc_lemmas:
        lemmas.update(str(l).upper() for l in doc_lemmas if str(l).strip())
    elif doc_tokens:
        # Fallback: если леммы недоступны, используем сами токены.
        lemmas.update(str(t).upper() for t in doc_tokens if str(t).strip())

    STOPWORD_LEMMA_CACHE[cache_key] = lemmas
    return lemmas


def _is_in_stopwords(ent_type: str, original_value: str) -> bool:
    """
    Проверяет, является ли распознанная сущность стоп-словом.

    Логика:
    1) сначала точное сравнение строки с элементами из ENTITY_STOPWORDS;
    2) затем сравнение по леммам, чтобы одно слово в стоп-списке
       отбрасывало все его формы.
    """
    ent_type_u = ent_type.upper()
    stopwords = ENTITY_STOPWORDS.get(ent_type_u)
    if not stopwords:
        return False

    value_u = original_value.upper().strip()
    if not value_u:
        return False

    # Точное совпадение по строке.
    if value_u in stopwords.exact:
        return True

    # Совпадение по regex.
    for rx in stopwords.regex:
        if rx.search(original_value):
            return True

    # Сравнение по леммам.
    value_lemmas = _get_lemmas_for_text(original_value)
    if not value_lemmas:
        return False

    for sw in stopwords.exact:
        sw_lemmas = _get_lemmas_for_text(sw)
        if value_lemmas & sw_lemmas:
            return True

    return False


def build_nlp_engine(language: str) -> NlpEngine:
    """
    Создаёт NLP-движок для Presidio на базе SpaCy.
    Для русского языка предполагается установленная модель ru_core_news_md/ru_core_news_lg.
    """
    # Можно переопределить через переменные окружения или конфиг при необходимости.
    nlp_configuration = {
        "nlp_engine_name": "spacy",
        "models": [
            {"lang_code": language, "model_name": f"{language}_core_news_md"},
        ],
    }
    provider = NlpEngineProvider(nlp_configuration=nlp_configuration)
    return provider.create_engine()


def _debug_log(msg: str, data: dict) -> None:
    # #region agent log
    try:
        log_path = Path(__file__).resolve().parent / ".cursor" / "debug.log"
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with open(log_path, "a", encoding="utf-8") as f:
            import json, time
            f.write(json.dumps({"message": msg, "data": data, "timestamp": int(time.time() * 1000), "runId": "post-fix"}, ensure_ascii=False) + "\n")
    except Exception:
        pass
    # #endregion

def build_analyzer(config: MaskingConfig) -> AnalyzerEngine:
    """
    Создаёт AnalyzerEngine Presidio, регистрируя кастомные recognizers для РФ.
    """
    # #region agent log
    _debug_log("build_analyzer entry", {"config_language": config.language, "hypothesisId": "C"})
    # #endregion
    nlp_engine = build_nlp_engine(config.language)
    registry = RecognizerRegistry(supported_languages=[config.language])
    # #region agent log
    _debug_log("after RecognizerRegistry()", {"registry_supported_languages": getattr(registry, "supported_languages", "N/A"), "hypothesisId": "A"})
    # #endregion
    registry.load_predefined_recognizers()
    # #region agent log
    _debug_log("before AnalyzerEngine", {"registry_supported_languages": getattr(registry, "supported_languages", "N/A"), "analyzer_langs": [config.language], "hypothesisId": "B"})
    # #endregion

    # Кастомные recognizers для РФ.

    # ИНН (10 или 12 цифр).
    inn_pattern = Pattern(name="inn_ru", regex=r"\b\d{10}(\d{2})?\b", score=0.6)
    inn_recognizer = PatternRecognizer(
        supported_entity="INN",
        patterns=[inn_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(inn_recognizer)

    # ОГРН (13 цифр).
    ogrn_pattern = Pattern(name="ogrn_ru", regex=r"\b\d{13}\b", score=0.6)
    ogrn_recognizer = PatternRecognizer(
        supported_entity="OGRN",
        patterns=[ogrn_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(ogrn_recognizer)

    # СНИЛС (формат 000-000-000 00).
    snils_pattern = Pattern(
        name="snils_ru",
        regex=r"\b\d{3}-\d{3}-\d{3}\s\d{2}\b",
        score=0.7,
    )
    snils_recognizer = PatternRecognizer(
        supported_entity="SNILS",
        patterns=[snils_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(snils_recognizer)

    # Банковский счёт (простейший шаблон для 20-значного счёта РФ).
    bank_acc_pattern = Pattern(
        name="bank_account_ru",
        regex=r"\b\d{20}\b",
        score=0.5,
    )
    bank_acc_recognizer = PatternRecognizer(
        supported_entity="BANK_ACCOUNT",
        patterns=[bank_acc_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(bank_acc_recognizer)

    # ИНН/КПП/ОГРН организаций как отдельные типы (примеры, можно доработать).
    org_inn_pattern = Pattern(
        name="org_inn_ru",
        regex=r"\b\d{10}\b",
        score=0.55,
    )
    org_inn_recognizer = PatternRecognizer(
        supported_entity="ORG_INN",
        patterns=[org_inn_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(org_inn_recognizer)

    org_kpp_pattern = Pattern(
        name="org_kpp_ru",
        regex=r"\b\d{9}\b",
        score=0.55,
    )
    org_kpp_recognizer = PatternRecognizer(
        supported_entity="ORG_KPP",
        patterns=[org_kpp_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(org_kpp_recognizer)

    org_ogrn_pattern = Pattern(
        name="org_ogrn_ru",
        regex=r"\b\d{13}\b",
        score=0.55,
    )
    org_ogrn_recognizer = PatternRecognizer(
        supported_entity="ORG_OGRN",
        patterns=[org_ogrn_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(org_ogrn_recognizer)

    # Дополнительные recognizers для e‑mail, телефонов и ИНН/КПП с явными метками.
    email_pattern = Pattern(
        name="email_ru_simple",
        regex=r"\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9.-]+\b",
        score=0.8,
    )
    email_recognizer = PatternRecognizer(
        supported_entity="EMAIL_ADDRESS",
        patterns=[email_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(email_recognizer)

    # Полный почтовый адрес РФ как единая LOCATION.
    # Пример: "640004, Курганская область, г. Курган, ул. Панфилова 22"
    # Логика: 6‑значный индекс, затем запятая и остальная часть строки до конца
    # строки или до точки с запятой.
    full_address_pattern = Pattern(
        name="ru_full_postal_address",
        regex=r"\b\d{6},[^\n;]+",
        score=0.9,
    )
    full_address_recognizer = PatternRecognizer(
        supported_entity="LOCATION",
        patterns=[full_address_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(full_address_recognizer)

    # ФИО в формате инициалов и фамилии: "Д.И. Анучин", "О.П. Чувардин",
    # а также варианты со пробелами между инициалами: "Д. И. Анучин".
    initials_person_pattern = Pattern(
        name="person_ru_initials_surname",
        regex=r"\b[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?",
        score=0.8,
    )
    initials_person_recognizer = PatternRecognizer(
        supported_entity="PERSON",
        patterns=[initials_person_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(initials_person_recognizer)

    phone_patterns = [
        Pattern(
            name="phone_ru_international",
            regex=r"(?:\+7|8)\s*(?:\(\d{3}\)|\d{3})[\s-]?\d{3}[\s-]?\d{2}[\s-]?\d{2}",
            score=0.8,
        ),
        Pattern(
            name="phone_ru_generic",
            regex=r"\b\+?\d[\d\-\s()]{9,}\d\b",
            score=0.5,
        ),
    ]
    phone_recognizer = PatternRecognizer(
        supported_entity="PHONE_NUMBER",
        patterns=phone_patterns,
        supported_language=config.language,
    )
    registry.add_recognizer(phone_recognizer)

    org_inn_labeled_pattern = Pattern(
        name="org_inn_labeled_ru",
        regex=r"\bИНН\s+(?:\d{10}|\d{12})\b",
        score=0.8,
    )
    org_inn_labeled_recognizer = PatternRecognizer(
        supported_entity="ORG_INN",
        patterns=[org_inn_labeled_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(org_inn_labeled_recognizer)

    org_kpp_labeled_pattern = Pattern(
        name="org_kpp_labeled_ru",
        regex=r"\bКПП\s+\d{9}\b",
        score=0.8,
    )
    org_kpp_labeled_recognizer = PatternRecognizer(
        supported_entity="ORG_KPP",
        patterns=[org_kpp_labeled_pattern],
        supported_language=config.language,
    )
    registry.add_recognizer(org_kpp_labeled_recognizer)

    analyzer = AnalyzerEngine(
        nlp_engine=nlp_engine,
        registry=registry,
        supported_languages=[config.language],
    )
    # Инициализируем NLP-движок для морфологической фильтрации стоп-слов.
    global STOPWORDS_NLP_ENGINE, STOPWORDS_LANGUAGE, STOPWORD_LEMMA_CACHE
    STOPWORDS_NLP_ENGINE = nlp_engine
    STOPWORDS_LANGUAGE = config.language
    STOPWORD_LEMMA_CACHE.clear()
    return analyzer


def get_files(input_dir: Path) -> Iterable[Path]:
    """
    Рекурсивный обход входной директории с фильтрацией по поддерживаемым расширениям.
    """
    for root, _, files in os.walk(input_dir):
        root_path = Path(root)
        for name in files:
            ext = Path(name).suffix.lower()
            if ext in SUPPORTED_EXTENSIONS:
                yield root_path / name


def _load_ocr_deps() -> Tuple["Image.Image", "PaddleOCR", "np"]:  # type: ignore[name-defined]
    global _OCR_DEPS
    if _OCR_DEPS is not None:
        return _OCR_DEPS

    # Отключаем проверки источников моделей и OneDNN (частая причина ошибок на CPU).
    os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
    os.environ.setdefault("FLAGS_use_mkldnn", "0")
    os.environ.setdefault("PADDLE_DISABLE_MKLDNN", "1")
    os.environ.setdefault("FLAGS_use_onednn", "0")
    os.environ.setdefault("FLAGS_use_new_executor", "0")
    os.environ.setdefault("FLAGS_enable_pir_api", "0")
    os.environ.setdefault("FLAGS_enable_pir_in_executor", "0")

    try:
        from paddleocr import PaddleOCR  # type: ignore
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Для OCR требуется пакет paddleocr. Ошибка: %s", exc)
        raise

    try:
        from PIL import Image  # type: ignore
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Для OCR требуется пакет Pillow. Ошибка: %s", exc)
        raise

    try:
        import numpy as np  # type: ignore
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Для OCR требуется пакет numpy. Ошибка: %s", exc)
        raise

    _OCR_DEPS = (Image, PaddleOCR, np)
    return _OCR_DEPS


def _get_ocr_engine() -> "PaddleOCR":  # type: ignore[name-defined]
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE

    _, PaddleOCR, _ = _load_ocr_deps()
    try:
        _OCR_ENGINE = PaddleOCR(
            lang=OCR_LANGUAGE,
            use_textline_orientation=True,
            device="cpu",
            enable_mkldnn=False,
        )
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Не удалось инициализировать PaddleOCR: %s", exc)
        raise
    return _OCR_ENGINE


def _ocr_image(path: Path, lang: str = OCR_LANGUAGE) -> str:
    Image, _, np = _load_ocr_deps()
    ocr = _get_ocr_engine()

    try:
        with Image.open(path) as img:
            frames = []
            try:
                from PIL import ImageSequence  # type: ignore

                frames = list(ImageSequence.Iterator(img))
            except Exception:
                frames = [img]

            texts: List[str] = []
            for frame in frames:
                frame_rgb = frame if frame.mode == "RGB" else frame.convert("RGB")
                result = ocr.predict(
                    np.array(frame_rgb),
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                    use_textline_orientation=False,
                )
                texts.append(_format_paddle_ocr(result))
            return "\n".join(t for t in texts if t.strip())
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Ошибка OCR для изображения %s: %s", path, exc)
        return ""


def _pdf_contains_images(path: Path) -> bool:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Для OCR PDF требуется PyMuPDF. Ошибка: %s", exc)
        return False

    try:
        with fitz.open(str(path)) as doc:
            for page in doc:
                if page.get_images(full=True):
                    return True
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Не удалось проверить PDF на наличие изображений %s: %s", path, exc)
    return False


def _ocr_pdf_images(path: Path, lang: str = OCR_LANGUAGE) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Для OCR PDF требуется PyMuPDF. Ошибка: %s", exc)
        return ""

    Image, _, np = _load_ocr_deps()
    ocr = _get_ocr_engine()

    texts: List[str] = []
    try:
        with fitz.open(str(path)) as doc:
            matrix = fitz.Matrix(OCR_PDF_ZOOM, OCR_PDF_ZOOM)
            for page_index, page in enumerate(doc, start=1):
                if not page.get_images(full=True):
                    continue
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                mode = "RGB" if pix.n < 4 else "RGBA"
                img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                if img.mode != "RGB":
                    img = img.convert("RGB")
                result = ocr.predict(
                    np.array(img),
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                    use_textline_orientation=False,
                )
                page_text = _format_paddle_ocr(result)
                if page_text.strip():
                    texts.append(f"--- PAGE {page_index} ---\n{page_text}")
    except Exception as exc:  # noqa: BLE001
        LOGGER.error("Ошибка OCR для PDF %s: %s", path, exc)
        return ""

    return "\n".join(texts)


def _format_paddle_ocr(result: List) -> str:
    if not result:
        return ""

    lines: List[str] = []
    for line in result:
        if isinstance(line, dict) and "rec_texts" in line:
            texts = line.get("rec_texts") or []
            scores = line.get("rec_scores") or []
            for idx, text_item in enumerate(texts):
                text_val = text_item[0] if isinstance(text_item, (list, tuple)) else text_item
                score_val = None
                if idx < len(scores):
                    try:
                        score_val = float(scores[idx])
                    except Exception:
                        score_val = None
                if score_val is not None and score_val < OCR_MIN_CONFIDENCE:
                    continue
                if text_val:
                    lines.append(str(text_val))
            continue

        if not line or len(line) < 2:
            continue
        text_data = line[1]
        if not text_data or len(text_data) < 2:
            continue
        text, score = text_data[0], text_data[1]
        if score is not None and score < OCR_MIN_CONFIDENCE:
            continue
        if text:
            lines.append(str(text))
    return "\n".join(lines)


def _build_ocr_output_path(output_dir: Path, rel_path: Path) -> Path:
    output_path = output_dir / rel_path
    return output_path.with_suffix(".ocr.txt")


def read_text_from_file(path: Path) -> str:
    """
    Чтение текста из файла в зависимости от расширения.
    Для .doc используется textract (при наличии), иначе логируем ошибку.
    """
    suffix = path.suffix.lower()

    if suffix == ".txt":
        with path.open("r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if suffix == ".docx":
        doc = Document(str(path))
        paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
        tables_text = ""
        if doc.tables:
            # Собираем текст из таблиц только для статистики (пока не используем при маскировании).
            tables_text = " ".join(
                cell.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
        # #region agent log
        _debug_log(
            "docx read stats",
            {
                "file_name": path.name,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "paragraphs_text_len": len(paragraphs_text),
                "tables_text_len": len(tables_text),
                "hypothesisId": "H1",
            },
        )
        # #endregion
        return paragraphs_text

    if suffix == ".pdf":
        # Поддержка текстовых PDF (изображения/сканы не распознаются).
        # Сначала пробуем PyMuPDF (лучше работает с нестандартными кодировками),
        # затем fallback на pdfminer.six.
        try:
            import fitz  # PyMuPDF

            with fitz.open(str(path)) as doc:
                pages_text = [page.get_text("text") for page in doc]
            return "\n".join(pages_text)
        except Exception as pymupdf_exc:  # noqa: BLE001
            LOGGER.debug("PyMuPDF не смог извлечь текст из %s: %s", path, pymupdf_exc)

        try:
            # Используем pdfminer.six (устанавливается как зависимость textract).
            from pdfminer.high_level import extract_text  # type: ignore
        except Exception as exc:  # noqa: BLE001
            # В старых версиях pdfminer.six (20181108) high_level.extract_text отсутствует.
            try:
                from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter  # type: ignore
                from pdfminer.converter import TextConverter  # type: ignore
                from pdfminer.layout import LAParams  # type: ignore
                from pdfminer.pdfpage import PDFPage  # type: ignore
            except Exception as inner_exc:  # noqa: BLE001
                LOGGER.error(
                    "Для обработки .pdf требуется pdfminer.six. Ошибка: %s; дополнительная ошибка: %s",
                    exc,
                    inner_exc,
                )
                raise

            def _extract_text_legacy(pdf_path: Path) -> str:
                output = StringIO()
                rsrcmgr = PDFResourceManager()
                laparams = LAParams()
                with TextConverter(rsrcmgr, output, laparams=laparams) as device:
                    interpreter = PDFPageInterpreter(rsrcmgr, device)
                    with pdf_path.open("rb") as f:
                        for page in PDFPage.get_pages(f):
                            interpreter.process_page(page)
                return output.getvalue()

            return _extract_text_legacy(path) or ""

        return extract_text(str(path)) or ""

    if suffix in IMAGE_EXTENSIONS:
        return _ocr_image(path, lang=OCR_LANGUAGE)

    if suffix == ".doc":
        try:
            import textract  # type: ignore
        except ImportError:
            LOGGER.error("Для обработки .doc требуется пакет textract. Файл %s будет пропущен.", path)
            raise

        text = textract.process(str(path))
        return text.decode("utf-8", errors="ignore")

    raise ValueError(f"Неподдерживаемое расширение файла: {path}")


def write_text_to_file(path: Path, text: str, original_suffix: str) -> None:
    """
    Запись маскированного текста в файл.
    Для .docx создаётся простой документ Word, для .txt — текстовый файл.
    Для .doc по умолчанию пишем .txt, сохраняя расширение .doc в имени.
    """
    suffix = original_suffix.lower()

    if suffix == ".txt" or suffix == ".doc":
        # Для .doc создаём текстовый файл с тем же именем.
        with path.open("w", encoding="utf-8") as f:
            f.write(text)
        return

    if suffix == ".docx":
        # #region agent log
        _debug_log(
            "docx write stats",
            {
                "file_name": path.name,
                "lines": len(text.splitlines()),
                "text_len": len(text),
                "hypothesisId": "H2",
            },
        )
        # #endregion
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(str(path))
        return

    if suffix == ".pdf":
        # Пишем простой текстовый PDF без сохранения исходной верстки.
        # Важно: для кириллицы нужен TTF-шрифт с поддержкой Unicode.
        path.parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        margin = 40
        y = height - margin
        line_height = 14
        max_chars = 110

        def _register_cyrillic_font() -> str:
            candidates = [
                Path(r"C:\Windows\Fonts\arial.ttf"),
                Path(r"C:\Windows\Fonts\times.ttf"),
                Path(r"C:\Windows\Fonts\calibri.ttf"),
                Path(r"C:\Windows\Fonts\verdana.ttf"),
                Path(r"C:\Windows\Fonts\tahoma.ttf"),
                Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
                Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
            ]

            for font_path in candidates:
                try:
                    if font_path.is_file():
                        font_name = f"ttf_{font_path.stem}"
                        if font_name not in pdfmetrics.getRegisteredFontNames():
                            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                        return font_name
                except Exception:
                    continue
            return "Helvetica"

        font_name = _register_cyrillic_font()
        c.setFont(font_name, 11)

        for raw_line in text.splitlines():
            wrapped = textwrap.wrap(raw_line, width=max_chars) or [""]
            for line in wrapped:
                if y - line_height < margin:
                    c.showPage()
                    c.setFont(font_name, 11)
                    y = height - margin
                c.drawString(margin, y, line)
                y -= line_height
        c.save()
        return

    # На всякий случай, как текст.
    with path.open("w", encoding="utf-8") as f:
        f.write(text)


class TokenGenerator:
    """
    Генератор токенов вида <TYPE_XXXX>, детерминированный в рамках запуска.
    """

    def __init__(self) -> None:
        self._counters: Dict[str, int] = defaultdict(int)
        self._mapping: Dict[Tuple[str, str], str] = {}

    @property
    def mapping(self) -> Dict[str, Dict[str, str]]:
        """
        Возвращает словарь для сохранения в JSON:
        {
          "<PERSON_0001>": {"type": "PERSON", "original": "Иванов Иван Иванович"},
          ...
        }
        """
        result: Dict[str, Dict[str, str]] = {}
        for (ent_type, original), token in self._mapping.items():
            result[token] = {"type": ent_type, "original": original}
        return result

    def get_token(self, ent_type: str, original: str) -> str:
        key = (ent_type, original)
        if key in self._mapping:
            return self._mapping[key]

        self._counters[ent_type] += 1
        num = self._counters[ent_type]
        token = f"<{ent_type.upper()}_{num:04d}>"
        self._mapping[key] = token
        return token


def mask_text(
    text: str,
    analyzer: AnalyzerEngine,
    config: MaskingConfig,
    token_gen: TokenGenerator,
) -> Tuple[str, List[RecognizerResult]]:
    """
    Запускает анализ Presidio и заменяет найденные сущности на токены.
    Возвращает маскированный текст и список исходных результатов.
    """
    results: List[RecognizerResult] = analyzer.analyze(
        text=text,
        language=config.language,
        entities=config.entity_types,
        score_threshold=config.score_threshold,
    )

    # Фильтрация результатов:
    # 1) не трогаем уже сгенерированные плейсхолдеры вида <TYPE_0001>, чтобы избежать
    #    "маскирования плейсхолдеров плейсхолдерами" и дублей в JSON;
    # 2) отбрасываем заведомые ложные срабатывания по словарю ENTITY_STOPWORDS.
    placeholder_re = re.compile(r"^<?[A-Z_]+_\d{4}>?$")

    filtered_results: List[RecognizerResult] = []
    for res in results:
        original_value_raw = text[res.start : res.end]
        original_value = original_value_raw.strip()

        # Пропускаем, если это уже плейсхолдер.
        if placeholder_re.match(original_value):
            continue

        # Пропускаем заведомо неверные сущности по словарю стоп-слов,
        # в том числе с учётом всех их склонений (через лемматизацию).
        if _is_in_stopwords(res.entity_type, original_value):
            continue

        filtered_results.append(res)

    # Сортируем по убыванию start, чтобы не смещать индексы при подстановке.
    results_sorted = sorted(filtered_results, key=lambda r: r.start, reverse=True)

    masked_text = text
    for res in results_sorted:
        original_value = masked_text[res.start : res.end]
        token = token_gen.get_token(res.entity_type, original_value)

        masked_text = masked_text[: res.start] + token + masked_text[res.end :]

    return masked_text, filtered_results


def process_docx_file(
    file_path: Path,
    output_file_path: Path,
    analyzer: AnalyzerEngine,
    config: MaskingConfig,
    token_gen: TokenGenerator,
    stats: Dict[str, int],
) -> None:
    """
    Обработка .docx-файла с сохранением структуры документа (таблицы, абзацы).
    Текст маскируется по абзацам и ячейкам таблиц.
    """
    try:
        doc = Document(str(file_path))
    except Exception as e:  # noqa: BLE001
        LOGGER.error("Ошибка открытия .docx файла %s: %s", file_path, e)
        stats["errors"] += 1
        return

    paragraphs = list(doc.paragraphs)
    table_paragraphs = [
        p
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
    ]

    total_par_len = sum(len(p.text) for p in paragraphs)
    total_table_len = sum(len(p.text) for p in table_paragraphs)

    # #region agent log
    _debug_log(
        "docx process before",
        {
            "file_name": file_path.name,
            "paragraphs": len(paragraphs),
            "table_paragraphs": len(table_paragraphs),
            "paragraphs_text_len": total_par_len,
            "tables_text_len": total_table_len,
            "hypothesisId": "H3",
        },
    )
    # #endregion

    all_results: List[RecognizerResult] = []

    def _mask_paragraphs(par_list: List["docx.text.paragraph.Paragraph"]) -> None:  # type: ignore[name-defined]
        nonlocal all_results
        for para in par_list:
            text = para.text
            if not text:
                continue
            masked_text, results = mask_text(
                text=text,
                analyzer=analyzer,
                config=config,
                token_gen=token_gen,
            )
            para.text = masked_text
            all_results.extend(results)

    _mask_paragraphs(paragraphs)
    _mask_paragraphs(table_paragraphs)

    total_par_len_after = sum(len(p.text) for p in paragraphs)
    total_table_len_after = sum(len(p.text) for p in table_paragraphs)

    # #region agent log
    _debug_log(
        "docx process after",
        {
            "file_name": file_path.name,
            "paragraphs_text_len_after": total_par_len_after,
            "tables_text_len_after": total_table_len_after,
            "entities_found": len(all_results),
            "hypothesisId": "H4",
        },
    )
    # #endregion

    # Обновление статистики по типам сущностей.
    for res in all_results:
        stats[res.entity_type] += 1

    try:
        LOGGER.debug("Сохранение .docx файла %s", output_file_path)
        output_file_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file_path))
        stats["files_processed"] += 1
    except Exception as e:  # noqa: BLE001
        LOGGER.error("Ошибка записи .docx файла %s: %s", output_file_path, e)
        stats["errors"] += 1


def process_file(
    file_path: Path,
    input_dir: Path,
    output_dir: Path,
    analyzer: AnalyzerEngine,
    config: MaskingConfig,
    token_gen: TokenGenerator,
    stats: Dict[str, int],
    enable_ocr: bool = False,
) -> None:
    """
    Обработка одного файла: чтение, маскирование, запись результата.
    Для .docx используется специальный путь с сохранением структуры документа.
    """
    try:
        rel_path = file_path.relative_to(input_dir)
    except ValueError:
        # На случай, если файл не лежит строго внутри input_dir.
        rel_path = file_path.name

    output_file_path = output_dir / rel_path
    output_file_path.parent.mkdir(parents=True, exist_ok=True)

    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        process_docx_file(
            file_path=file_path,
            output_file_path=output_file_path,
            analyzer=analyzer,
            config=config,
            token_gen=token_gen,
            stats=stats,
        )
        return

    if suffix in IMAGE_EXTENSIONS:
        if not enable_ocr:
            LOGGER.info("Пропуск изображения (OCR выключен): %s", file_path)
            return
        try:
            text = read_text_from_file(file_path)
        except Exception as e:  # noqa: BLE001
            LOGGER.error("Ошибка OCR изображения %s: %s", file_path, e)
            stats["errors"] += 1
            return

        try:
            masked_text, results = mask_text(
                text=text,
                analyzer=analyzer,
                config=config,
                token_gen=token_gen,
            )
        except Exception as e:  # noqa: BLE001
            LOGGER.error("Ошибка анализа OCR текста %s: %s", file_path, e)
            stats["errors"] += 1
            return

        for res in results:
            stats[res.entity_type] += 1

        try:
            ocr_output_path = _build_ocr_output_path(output_dir, rel_path)
            LOGGER.debug("Запись OCR результата %s", ocr_output_path)
            ocr_output_path.parent.mkdir(parents=True, exist_ok=True)
            ocr_output_path.write_text(masked_text, encoding="utf-8")
            stats["files_processed"] += 1
        except Exception as e:  # noqa: BLE001
            LOGGER.error("Ошибка записи OCR результата %s: %s", file_path, e)
            stats["errors"] += 1
        return

    try:
        LOGGER.debug("Чтение файла %s", file_path)
        text = read_text_from_file(file_path)
    except Exception as e:  # noqa: BLE001
        LOGGER.error("Ошибка чтения файла %s: %s", file_path, e)
        stats["errors"] += 1
        return

    try:
        masked_text, results = mask_text(
            text=text,
            analyzer=analyzer,
            config=config,
            token_gen=token_gen,
        )
    except Exception as e:  # noqa: BLE001
        LOGGER.error("Ошибка анализа файла %s: %s", file_path, e)
        stats["errors"] += 1
        return

    # Обновление статистики по типам сущностей.
    for res in results:
        stats[res.entity_type] += 1

    try:
        LOGGER.debug("Запись маскированного файла %s", output_file_path)
        write_text_to_file(output_file_path, masked_text, file_path.suffix)
        stats["files_processed"] += 1
    except Exception as e:  # noqa: BLE001
        LOGGER.error("Ошибка записи файла %s: %s", output_file_path, e)
        stats["errors"] += 1
        return

    if enable_ocr and suffix == ".pdf":
        if _pdf_contains_images(file_path):
            ocr_text = _ocr_pdf_images(file_path, lang=OCR_LANGUAGE)
            if ocr_text.strip():
                try:
                    ocr_masked, ocr_results = mask_text(
                        text=ocr_text,
                        analyzer=analyzer,
                        config=config,
                        token_gen=token_gen,
                    )
                    for res in ocr_results:
                        stats[res.entity_type] += 1
                    ocr_output_path = _build_ocr_output_path(output_dir, rel_path)
                    LOGGER.debug("Запись OCR результата PDF %s", ocr_output_path)
                    ocr_output_path.parent.mkdir(parents=True, exist_ok=True)
                    ocr_output_path.write_text(ocr_masked, encoding="utf-8")
                except Exception as e:  # noqa: BLE001
                    LOGGER.error("Ошибка OCR/записи PDF %s: %s", file_path, e)
                    stats["errors"] += 1


def setup_logging(log_level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, log_level.upper(), logging.INFO),
        format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Утилита маскирования персональных и корпоративных данных в текстовых документах (Microsoft Presidio)."
    )
    parser.add_argument(
        "--input",
        default="./input_texts",
        help="Путь к входной папке с текстовыми файлами (doc, docx, txt, pdf) и изображениями. По умолчанию ./input_texts.",
    )
    parser.add_argument(
        "--output",
        default="./masked_texts",
        help="Путь к выходной папке для маскированных файлов. По умолчанию ./masked_texts.",
    )
    parser.add_argument(
        "--map",
        default="./deanonymization_map.json",
        help="Путь к JSON-файлу со словарём демаскирования. По умолчанию ./deanonymization_map.json.",
    )
    parser.add_argument(
        "--language",
        default=None,
        help="Язык текста (по умолчанию ru или значение из config.yaml).",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        help="Уровень логирования.",
    )
    parser.add_argument(
        "--config",
        default="config.yaml",
        help="Путь к YAML-конфигу с параметрами Presidio и списка сущностей.",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Включить OCR для изображений и сканированных PDF. По умолчанию OCR выключен.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    setup_logging(args.log_level)

    input_dir = Path(args.input).resolve()
    output_dir = Path(args.output).resolve()
    map_path = Path(args.map).resolve()
    config_path = Path(args.config).resolve() if args.config else None

    if not input_dir.is_dir():
        raise SystemExit(f"Входная директория не найдена: {input_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)

    config = load_config(config_path)
    if args.language:
        config.language = args.language

    # Загрузка стоп-слов из YAML (по умолчанию stopwords.yaml рядом с config.yaml).
    global ENTITY_STOPWORDS
    stopwords_path: Optional[Path] = None
    if config_path is not None:
        stopwords_path = config_path.with_name("stopwords.yaml")
    ENTITY_STOPWORDS = load_stopwords(stopwords_path)

    LOGGER.info("Загрузка Presidio Analyzer (язык=%s)...", config.language)
    analyzer = build_analyzer(config)

    token_gen = TokenGenerator()

    stats: Dict[str, int] = defaultdict(int)

    files = list(get_files(input_dir))
    LOGGER.info("Найдено файлов для обработки: %d", len(files))

    for file_path in files:
        LOGGER.debug("Обработка файла %s", file_path)
        try:
            process_file(
                file_path=file_path,
                input_dir=input_dir,
                output_dir=output_dir,
                analyzer=analyzer,
                config=config,
                token_gen=token_gen,
                stats=stats,
                enable_ocr=args.ocr,
            )
        except Exception as e:  # noqa: BLE001
            LOGGER.error("Необработанное исключение при обработке файла %s: %s", file_path, e)
            stats["errors"] += 1

    # Сохранение словаря демаскирования.
    LOGGER.info("Сохранение словаря демаскирования в %s", map_path)
    with map_path.open("w", encoding="utf-8") as f:
        json.dump(token_gen.mapping, f, ensure_ascii=False, indent=2)

    # Формирование отчёта.
    LOGGER.info("Обработка завершена.")
    LOGGER.info("Файлов успешно обработано: %d", stats.get("files_processed", 0))
    LOGGER.info("Ошибок: %d", stats.get("errors", 0))

    LOGGER.info("Найденные сущности по типам:")
    for key, value in sorted(stats.items()):
        if key in {"files_processed", "errors"}:
            continue
        LOGGER.info("  %s: %d", key, value)


if __name__ == "__main__":
    main()
